# /utils/documents.py
"""Shared document text extraction and chunking.

Used by both the file manager (summarising a single document) and the
knowledge base (indexing whole folders for retrieval-augmented answers).
Every extractor degrades to an empty string rather than raising.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List, Tuple

from utils.helpers import clean_text, read_text_file
from utils.logger import get_logger

logger = get_logger("utils.documents")

# Extensions we can pull text out of.
TEXT_LIKE = {
    ".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".tsv", ".json", ".yaml",
    ".yml", ".toml", ".ini", ".cfg", ".py", ".js", ".ts", ".tsx", ".jsx", ".java",
    ".c", ".h", ".cpp", ".go", ".rs", ".rb", ".php", ".sh", ".sql", ".tex",
}
RICH_FORMATS = {".pdf", ".docx", ".docm", ".epub", ".html", ".htm", ".xml", ".pptx"}
SUPPORTED = TEXT_LIKE | RICH_FORMATS


def is_supported(path: str | Path) -> bool:
    """Return True when :func:`extract_text` can handle this file type."""
    return Path(path).suffix.lower() in SUPPORTED


def extract_text(path: str | Path, limit: int = 400_000) -> str:
    """Extract plain text from a document.

    Supports PDF, DOCX, PPTX, EPUB/HTML/XML and any plain-text format. Returns
    ``""`` when the file cannot be read or the optional parser is missing.

    Args:
        path: File to read.
        limit: Maximum number of characters to return.

    Returns:
        The extracted text.
    """
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if not file_path.exists() or not file_path.is_file():
        return ""

    try:
        if suffix == ".pdf":
            return _extract_pdf(file_path, limit)
        if suffix in {".docx", ".docm"}:
            return _extract_docx(file_path, limit)
        if suffix == ".pptx":
            return _extract_pptx(file_path, limit)
        if suffix in {".epub", ".html", ".htm", ".xml"}:
            raw = read_text_file(file_path, limit * 2)
            return clean_text(re.sub(r"<[^>]+>", " ", raw))[:limit]
        return read_text_file(file_path, limit)
    except Exception as exc:
        logger.debug("Extraction failed for %s: %s", file_path, exc)
        return ""


def _extract_pdf(path: Path, limit: int) -> str:
    """Pull text out of a PDF with pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages: List[str] = []
        total = 0
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
            total += len(text)
            if total >= limit:
                break
        return clean_text("\n".join(pages))[:limit]
    except Exception as exc:
        logger.debug("pypdf failed on %s: %s", path.name, exc)
        return ""


def _extract_docx(path: Path, limit: int) -> str:
    """Pull text (including tables) out of a DOCX with python-docx."""
    try:
        import docx

        document = docx.Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return clean_text("\n".join(parts))[:limit]
    except Exception as exc:
        logger.debug("python-docx failed on %s: %s", path.name, exc)
        return ""


def _extract_pptx(path: Path, limit: int) -> str:
    """Pull slide text out of a PPTX (zip + XML, no extra dependency)."""
    try:
        import zipfile

        parts: List[str] = []
        with zipfile.ZipFile(path) as archive:
            slides = sorted(
                name for name in archive.namelist()
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            )
            for slide in slides:
                xml = archive.read(slide).decode("utf-8", errors="replace")
                parts.extend(re.findall(r"<a:t>([^<]*)</a:t>", xml))
        return clean_text("\n".join(parts))[:limit]
    except Exception as exc:
        logger.debug("pptx extraction failed on %s: %s", path.name, exc)
        return ""


def chunk_text(
    text: str, chunk_size: int = 1200, overlap: int = 150
) -> List[Tuple[int, str]]:
    """Split text into overlapping chunks on paragraph/sentence boundaries.

    Args:
        text: The document text.
        chunk_size: Target characters per chunk.
        overlap: Characters of context repeated between chunks.

    Returns:
        A list of ``(index, chunk)`` tuples.
    """
    cleaned = clean_text(text)
    if not cleaned:
        return []
    if len(cleaned) <= chunk_size:
        return [(0, cleaned)]

    chunks: List[Tuple[int, str]] = []
    paragraphs = re.split(r"\n{2,}", cleaned)
    current = ""
    index = 0

    def flush(buffer: str) -> str:
        """Emit a chunk and return the overlap tail for the next one."""
        nonlocal index
        buffer = buffer.strip()
        if not buffer:
            return ""
        chunks.append((index, buffer))
        index += 1
        return buffer[-overlap:] if overlap > 0 else ""

    for paragraph in paragraphs:
        while len(paragraph) > chunk_size:
            # A single huge paragraph: cut it on sentence boundaries.
            cut = paragraph.rfind(". ", 0, chunk_size)
            cut = cut + 1 if cut > chunk_size // 2 else chunk_size
            current = flush(current + "\n\n" + paragraph[:cut])
            paragraph = paragraph[cut:].lstrip()
        if len(current) + len(paragraph) + 2 > chunk_size:
            current = flush(current)
        current = f"{current}\n\n{paragraph}".strip() if current else paragraph
    flush(current)
    return chunks


def document_metadata(path: Path) -> Dict[str, str]:
    """Return simple metadata describing a document."""
    try:
        stat = path.stat()
        return {
            "path": str(path),
            "name": path.name,
            "suffix": path.suffix.lower(),
            "size": str(stat.st_size),
            "modified": str(int(stat.st_mtime)),
        }
    except Exception:
        return {"path": str(path), "name": path.name, "suffix": path.suffix.lower()}


__all__ = ["extract_text", "chunk_text", "is_supported", "document_metadata", "SUPPORTED"]
